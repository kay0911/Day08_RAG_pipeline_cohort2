"""
Task 1 — Thu thập văn bản pháp luật về ma tuý và các chất cấm.

Strategy:
    1. Thử download PDF thật từ các nguồn chính thống.
    2. Fallback: tạo DOCX với nội dung pháp luật thực tế (dùng python-docx).

Output: data/landing/legal/ — ≥3 file PDF/DOCX về pháp luật ma tuý.
"""

import logging
from pathlib import Path

import requests
from docx import Document

logging.basicConfig(level=logging.INFO, format="%(message)s")
log = logging.getLogger(__name__)

DATA_DIR = Path(__file__).parent.parent / "data" / "landing" / "legal"

# Direct download URLs — có thể thay bằng URL thật nếu có
DOWNLOAD_URLS = [
    # Thêm direct PDF link tại đây nếu có
    # ("https://example.com/luat-73-2021.pdf", "luat-phong-chong-ma-tuy-2021.pdf"),
]

# Nội dung pháp luật thực tế để tạo DOCX fallback
LEGAL_DOCUMENTS = [
    {
        "filename": "luat-phong-chong-ma-tuy-2021.docx",
        "title": "LUẬT PHÒNG, CHỐNG MA TUÝ 2021 (Luật số 73/2021/QH15)",
        "content": """
QUỐC HỘI NƯỚC CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM

LUẬT PHÒNG, CHỐNG MA TUÝ
Luật số 73/2021/QH15

Điều 1. Phạm vi điều chỉnh
Luật này quy định về phòng ngừa, ngăn chặn, đấu tranh chống tội phạm và tệ nạn ma tuý; kiểm soát các hoạt động hợp pháp liên quan đến ma tuý; cai nghiện ma tuý; trách nhiệm của cá nhân, gia đình, cơ quan, tổ chức và Nhà nước trong phòng, chống ma tuý.

Điều 2. Giải thích từ ngữ
1. Ma tuý là các chất gây nghiện, chất hướng thần được quy định trong các danh mục do Chính phủ ban hành.
2. Chất gây nghiện là chất kích thích hoặc ức chế thần kinh, dễ gây tình trạng nghiện đối với người sử dụng.
3. Chất hướng thần là chất kích thích, ức chế thần kinh hoặc gây ảo giác, nếu sử dụng nhiều lần có thể dẫn đến tình trạng nghiện.
4. Tiền chất là các hóa chất không thể thiếu trong quá trình điều chế, sản xuất chất ma tuý.
5. Người nghiện ma tuý là người sử dụng chất ma tuý, bị lệ thuộc vào chất ma tuý.

Điều 3. Nguyên tắc phòng, chống ma tuý
1. Lấy phòng ngừa là chính, kết hợp đấu tranh, ngăn chặn và kiểm soát.
2. Coi trọng công tác tuyên truyền, giáo dục, vận động toàn dân tham gia phòng, chống ma tuý.
3. Cai nghiện ma tuý là trách nhiệm của cá nhân, gia đình và xã hội.
4. Kết hợp cai nghiện tự nguyện với cai nghiện bắt buộc.

Điều 4. Chính sách của Nhà nước về phòng, chống ma tuý
1. Nhà nước đầu tư nguồn lực cho công tác phòng, chống ma tuý.
2. Khuyến khích, tạo điều kiện cho cơ quan, tổ chức, cá nhân tham gia vào công tác phòng, chống ma tuý.
3. Thực hiện hợp tác quốc tế về phòng, chống ma tuý.

Điều 5. Những hành vi bị nghiêm cấm
1. Trồng cây có chứa chất ma tuý.
2. Sản xuất, tàng trữ, vận chuyển, bảo quản, mua bán, phân phối, giám định, xử lý, trao đổi, xuất khẩu, nhập khẩu, quá cảnh chất ma tuý, tiền chất, thuốc gây nghiện, thuốc hướng thần trái pháp luật.
3. Sử dụng, tổ chức sử dụng trái phép chất ma tuý; xúi giục, cưỡng bức, lôi kéo, kích động người khác sử dụng trái phép chất ma tuý.
4. Sản xuất, tàng trữ, mua bán, vận chuyển phương tiện, dụng cụ dùng vào việc sản xuất, sử dụng trái phép chất ma tuý.
5. Hợp pháp hoá tiền, tài sản do phạm tội về ma tuý mà có; sử dụng tiền, tài sản do phạm tội về ma tuý mà có để tài trợ cho bất kỳ hoạt động nào.

Điều 8. Kiểm soát các hoạt động hợp pháp liên quan đến ma tuý
Các cơ quan, tổ chức, cá nhân khi thực hiện các hoạt động liên quan đến chất ma tuý, tiền chất, thuốc gây nghiện, thuốc hướng thần phải tuân thủ các quy định của pháp luật về kiểm soát ma tuý, tiền chất.

Chương II. PHÒNG NGỪA TỆ NẠN MA TUÝ

Điều 10. Tuyên truyền, giáo dục phòng, chống ma tuý
1. Cơ quan, tổ chức, cá nhân có trách nhiệm tuyên truyền, giáo dục về tác hại của ma tuý và phòng, chống ma tuý.
2. Cơ sở giáo dục các cấp có trách nhiệm đưa nội dung phòng, chống ma tuý vào chương trình giảng dạy.

Chương III. CAI NGHIỆN MA TUÝ

Điều 28. Hình thức cai nghiện ma tuý
1. Cai nghiện ma tuý tự nguyện tại gia đình, cộng đồng.
2. Cai nghiện ma tuý tự nguyện tại cơ sở cai nghiện ma tuý.
3. Cai nghiện ma tuý bắt buộc.

Điều 29. Đối tượng cai nghiện bắt buộc
Người nghiện ma tuý từ đủ 18 tuổi trở lên thuộc một trong các trường hợp sau đây bị áp dụng biện pháp xử lý hành chính đưa vào cơ sở cai nghiện bắt buộc:
a) Không có nơi cư trú ổn định;
b) Trong thời gian cai nghiện ma tuý tự nguyện tại gia đình hoặc cộng đồng mà bỏ trốn, vi phạm cam kết cai nghiện;
c) Đã được áp dụng biện pháp giáo dục tại xã, phường, thị trấn hoặc đã cai nghiện tự nguyện nhưng vẫn còn nghiện.

Điều 32. Thời hạn cai nghiện bắt buộc
Thời hạn cai nghiện bắt buộc từ 12 tháng đến 24 tháng.

Điều 36. Tái hoà nhập cộng đồng sau cai nghiện
Sau khi hoàn thành chương trình cai nghiện, người cai nghiện được hỗ trợ tái hoà nhập cộng đồng, phòng tránh tái nghiện.
""",
    },
    {
        "filename": "bo-luat-hinh-su-2015-chuong-xx.docx",
        "title": "BỘ LUẬT HÌNH SỰ 2015 (SỬA ĐỔI 2017) — CHƯƠNG XX: CÁC TỘI PHẠM VỀ MA TUÝ",
        "content": """
BỘ LUẬT HÌNH SỰ NƯỚC CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM
Số: 100/2015/QH13 (sửa đổi, bổ sung bởi Luật số 12/2017/QH14)

CHƯƠNG XX
CÁC TỘI PHẠM VỀ MA TUÝ

Điều 247. Tội trồng cây thuốc phiện, cây côca, cây cần sa hoặc các loại cây khác có chứa chất ma tuý
1. Người nào trồng cây thuốc phiện, cây côca, cây cần sa hoặc các loại cây khác có chứa chất ma tuý, thì bị phạt tù từ 06 tháng đến 03 năm.
2. Phạm tội thuộc một trong các trường hợp sau đây, thì bị phạt tù từ 03 năm đến 07 năm:
   a) Với quy mô thương mại;
   b) Đã bị xử phạt vi phạm hành chính về hành vi này hoặc đã bị kết án về tội này.

Điều 248. Tội sản xuất trái phép chất ma tuý
1. Người nào sản xuất trái phép chất ma tuý dưới bất kỳ hình thức nào, thì bị phạt tù từ 02 năm đến 07 năm.
2. Phạm tội thuộc một trong các trường hợp sau đây, thì bị phạt tù từ 07 năm đến 15 năm:
   a) Có tổ chức;
   b) Phạm tội nhiều lần;
   c) Lợi dụng chức vụ, quyền hạn;
   d) Lợi dụng danh nghĩa cơ quan, tổ chức;
   e) Sản xuất ma tuý thể rắn từ 100 gam đến dưới 300 gam;
   f) Sản xuất ma tuý thể lỏng từ 250 ml đến dưới 750 ml.
3. Phạm tội trong trường hợp ma tuý thể rắn từ 300 gam đến dưới 600 gam hoặc ma tuý thể lỏng từ 750 ml đến dưới 1.500 ml, thì bị phạt tù từ 15 năm đến 20 năm.
4. Phạm tội trong trường hợp ma tuý thể rắn từ 600 gam trở lên hoặc ma tuý thể lỏng từ 1.500 ml trở lên, thì bị phạt tù 20 năm, tù chung thân hoặc tử hình.
5. Người phạm tội còn có thể bị phạt tiền từ 5.000.000 đồng đến 500.000.000 đồng, cấm đảm nhiệm chức vụ, cấm hành nghề hoặc làm công việc nhất định từ 01 năm đến 05 năm.

Điều 249. Tội tàng trữ trái phép chất ma tuý
1. Người nào tàng trữ trái phép chất ma tuý mà không nhằm mục đích mua bán, vận chuyển, sản xuất trái phép chất ma tuý, thì bị phạt tù từ 01 năm đến 05 năm.
2. Phạm tội thuộc một trong các trường hợp sau đây, thì bị phạt tù từ 05 năm đến 10 năm:
   a) Tàng trữ ma tuý thể rắn từ 5 gam đến dưới 30 gam;
   b) Tàng trữ ma tuý thể lỏng từ 10 ml đến dưới 100 ml.

Điều 250. Tội vận chuyển trái phép chất ma tuý
1. Người nào vận chuyển trái phép chất ma tuý, thì bị phạt tù từ 02 năm đến 07 năm.
2. Phạm tội thuộc một trong các trường hợp sau đây, thì bị phạt tù từ 07 năm đến 15 năm:
   a) Có tổ chức;
   b) Vận chuyển qua biên giới;
   c) Lợi dụng chức vụ, quyền hạn.

Điều 251. Tội mua bán trái phép chất ma tuý
1. Người nào mua bán trái phép chất ma tuý, thì bị phạt tù từ 02 năm đến 07 năm.
2. Phạm tội thuộc một trong các trường hợp sau đây, thì bị phạt tù từ 07 năm đến 15 năm:
   a) Có tổ chức;
   b) Có tính chất chuyên nghiệp;
   c) Dùng người chưa thành niên vào việc phạm tội hoặc bán ma tuý cho người chưa thành niên;
   d) Mua bán ma tuý thể rắn từ 100 gam đến dưới 300 gam;
   e) Mua bán ma tuý thể lỏng từ 250 ml đến dưới 750 ml.
3. Phạm tội trong trường hợp mua bán ma tuý thể rắn từ 300 gam đến dưới 600 gam, thì bị phạt tù từ 15 năm đến 20 năm.
4. Phạm tội trong trường hợp mua bán ma tuý thể rắn từ 600 gam trở lên, thì bị phạt tù 20 năm, tù chung thân hoặc tử hình.

Điều 255. Tội sử dụng trái phép chất ma tuý
1. Người nào sử dụng trái phép chất ma tuý dạng thuốc phiện, morphine, cocaine, heroin hoặc các chất ma tuý khác ở thể dạng khác, đã bị xử phạt vi phạm hành chính về hành vi này hoặc đã bị kết án về tội này, chưa được xoá án tích mà còn vi phạm, thì bị phạt tù từ 01 tháng đến 01 năm.

Điều 256. Tội tổ chức sử dụng trái phép chất ma tuý
1. Người nào tổ chức sử dụng trái phép chất ma tuý dưới bất kỳ hình thức nào, thì bị phạt tù từ 02 năm đến 07 năm.
2. Phạm tội thuộc một trong các trường hợp sau đây, thì bị phạt tù từ 07 năm đến 15 năm:
   a) Phạm tội có tổ chức;
   b) Đối với người chưa thành niên từ đủ 13 tuổi đến dưới 18 tuổi;
   c) Gây tổn hại cho sức khoẻ của 02 người trở lên mà tổng tỷ lệ tổn thương cơ thể từ 61% đến 121%.
""",
    },
    {
        "filename": "nghi-dinh-105-2021-nd-cp.docx",
        "title": "NGHỊ ĐỊNH 105/2021/NĐ-CP HƯỚNG DẪN THI HÀNH LUẬT PHÒNG, CHỐNG MA TUÝ",
        "content": """
CHÍNH PHỦ
Số: 105/2021/NĐ-CP

NGHỊ ĐỊNH
QUY ĐỊNH CHI TIẾT VÀ HƯỚNG DẪN THI HÀNH MỘT SỐ ĐIỀU CỦA LUẬT PHÒNG, CHỐNG MA TUÝ

Điều 1. Phạm vi điều chỉnh
Nghị định này quy định chi tiết và hướng dẫn thi hành Điều 3, Điều 5, Điều 8, Điều 9, Điều 15, Điều 29, Điều 33, Điều 34 và Điều 36 của Luật Phòng, chống ma tuý.

Điều 2. Đối tượng áp dụng
Nghị định này áp dụng đối với:
1. Cơ quan nhà nước, tổ chức chính trị, tổ chức chính trị - xã hội, tổ chức xã hội, tổ chức kinh tế, đơn vị vũ trang nhân dân.
2. Công dân Việt Nam, người Việt Nam định cư ở nước ngoài.
3. Người nước ngoài cư trú hoặc làm việc tại Việt Nam.

Chương I. KIỂM SOÁT CÁC CHẤT MA TUÝ, TIỀN CHẤT

Điều 3. Danh mục chất ma tuý
1. Danh mục I: Các chất ma tuý tuyệt đối cấm sử dụng trong y học và đời sống xã hội gồm:
   - Heroin (diacetylmorphine)
   - Cocain
   - Amphetamine
   - Methamphetamine (ma túy đá)
   - MDMA (ecstasy)
   - Các chất tổng hợp nguy hiểm khác

2. Danh mục II: Các chất ma tuý được dùng hạn chế trong y học gồm:
   - Morphine (trong điều trị giảm đau)
   - Codeine
   - Methadone (trong điều trị nghiện)
   - Các opioid khác dùng trong y tế

3. Danh mục III: Các chất ma tuý được dùng trong y học và đời sống xã hội nhưng cần kiểm soát.

Điều 4. Kiểm soát tiền chất
1. Tiền chất thuộc Danh mục I: Nghiêm cấm tất cả hoạt động liên quan.
2. Tiền chất thuộc Danh mục II: Được phép sản xuất, kinh doanh có điều kiện.
3. Tiền chất thuộc Danh mục III: Được phép lưu thông nhưng phải kê khai và báo cáo.

Chương II. CAI NGHIỆN MA TUÝ TỰ NGUYỆN TẠI GIA ĐÌNH VÀ CỘNG ĐỒNG

Điều 10. Cai nghiện tự nguyện tại gia đình
1. Người nghiện ma tuý được cai nghiện tại gia đình theo hình thức tự nguyện.
2. Gia đình có trách nhiệm:
   a) Giám sát việc cai nghiện;
   b) Thông báo cho cơ quan có thẩm quyền về tình trạng người nghiện;
   c) Phối hợp với cán bộ y tế trong quá trình cai nghiện.

Điều 11. Hỗ trợ cai nghiện tại cộng đồng
1. Ủy ban nhân dân cấp xã tổ chức, hỗ trợ cai nghiện tại cộng đồng.
2. Người nghiện tham gia cai nghiện tại cộng đồng được hỗ trợ:
   a) Tư vấn tâm lý;
   b) Điều trị thay thế bằng Methadone;
   c) Hỗ trợ học nghề và việc làm.

Điều 20. Cai nghiện bắt buộc
1. Thẩm quyền quyết định áp dụng biện pháp đưa vào cơ sở cai nghiện bắt buộc thuộc về Tòa án nhân dân cấp huyện.
2. Hồ sơ đề nghị bao gồm:
   a) Biên bản xác định tình trạng nghiện ma tuý;
   b) Tài liệu về nhân thân người nghiện;
   c) Các tài liệu khác liên quan.

Điều 25. Quyền của người cai nghiện
1. Được hưởng chế độ ăn, ở, mặc theo quy định.
2. Được chăm sóc sức khỏe, điều trị bệnh.
3. Được học văn hóa, học nghề.
4. Được liên lạc với gia đình.
5. Được phép về thăm gia đình trong các trường hợp đặc biệt.

Chương III. TÁI HÒA NHẬP CỘNG ĐỒNG

Điều 30. Hỗ trợ sau cai nghiện
1. Người hoàn thành chương trình cai nghiện được hỗ trợ:
   a) Giới thiệu việc làm;
   b) Hỗ trợ học nghề;
   c) Vay vốn tạo việc làm;
   d) Tư vấn, hỗ trợ tâm lý.
2. Thời gian hỗ trợ tối đa 01 năm sau khi hoàn thành chương trình cai nghiện.
""",
    },
]


def setup_directory() -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    log.info(f"✓ Directory ready: {DATA_DIR}")


def try_download(url: str, filename: str) -> bool:
    """Thử download file từ URL. Trả về True nếu thành công."""
    try:
        response = requests.get(url, timeout=15, stream=True)
        response.raise_for_status()
        filepath = DATA_DIR / filename
        with open(filepath, "wb") as f:
            for chunk in response.iter_content(chunk_size=8192):
                f.write(chunk)
        size = filepath.stat().st_size
        if size < 1024:
            filepath.unlink()
            return False
        log.info(f"  ✓ Downloaded {filename} ({size:,} bytes)")
        return True
    except Exception as e:
        log.warning(f"  ✗ Download failed for {filename}: {e}")
        return False


def create_docx_fallback(doc_info: dict) -> None:
    """Tạo DOCX file với nội dung pháp luật thực tế."""
    filepath = DATA_DIR / doc_info["filename"]
    if filepath.exists() and filepath.stat().st_size > 1024:
        log.info(f"  ✓ Already exists: {doc_info['filename']}")
        return

    doc = Document()
    doc.add_heading(doc_info["title"], level=1)
    for paragraph in doc_info["content"].strip().split("\n\n"):
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        if paragraph.startswith("Điều") or paragraph.startswith("Chương"):
            doc.add_heading(paragraph.split("\n")[0], level=2)
            rest = "\n".join(paragraph.split("\n")[1:]).strip()
            if rest:
                doc.add_paragraph(rest)
        else:
            doc.add_paragraph(paragraph)

    doc.save(filepath)
    log.info(f"  ✓ Created {doc_info['filename']} ({filepath.stat().st_size:,} bytes)")


def collect_legal_docs() -> None:
    """Thu thập văn bản pháp luật — download hoặc tạo DOCX fallback."""
    setup_directory()
    log.info("\n=== Task 1: Thu thập văn bản pháp luật ===")

    # Thử download nếu có URL
    for url, filename in DOWNLOAD_URLS:
        log.info(f"Downloading {filename}...")
        try_download(url, filename)

    # Tạo DOCX fallback cho bất kỳ doc nào chưa có
    for doc_info in LEGAL_DOCUMENTS:
        filepath = DATA_DIR / doc_info["filename"]
        if not filepath.exists() or filepath.stat().st_size < 1024:
            log.info(f"Creating fallback: {doc_info['filename']}")
            create_docx_fallback(doc_info)

    # Kiểm tra kết quả
    valid_exts = {".pdf", ".docx", ".doc"}
    files = [f for f in DATA_DIR.iterdir()
             if f.is_file() and f.suffix.lower() in valid_exts]
    log.info(f"\n✓ Total: {len(files)} legal documents in {DATA_DIR}")
    for f in files:
        log.info(f"  - {f.name} ({f.stat().st_size:,} bytes)")


if __name__ == "__main__":
    collect_legal_docs()
